import docx
import logging
import os
from data_loaders.data_loader_interface import DataLoaderInterface
from msuliot.base_64 import Base64

class DocxLoader(DataLoaderInterface):


    def load_data(self, data_path):
        json_doc = self.convert_docx_to_json(data_path)
        return " ".join(item['text'] for item in json_doc['content'])


    def save_image(self, image, image_path):
        """
        Save a DOCX image to a file.
        """
        with open(image_path, 'wb') as f:
            f.write(image.blob)


    def convert_docx_to_json(self, filepath, image_dir='images'):
        try:
            doc = docx.Document(filepath)
        except Exception as e:
            logging.error(f"Error opening DOCX file {filepath}: {e}")
            return {}
        
        # if not os.path.exists(image_dir):
        #     os.makedirs(image_dir)

        images_info = []
        
        # for rel_key, rel in doc.part.rels.items():

        #     if "image" in rel.target_ref:
        #         image_part = rel.target_part
        #         image_ext = image_part.content_type.split('/')[1]
        #         image_filename = f'image_{rel_key}.{image_ext}'
        #         image_path = os.path.join(image_dir, image_filename)
        
        #         self.save_image(image_part, image_path)
                
        #         images_info.append({
        #             "filename": image_filename,
        #             "path": image_path
        #         })

        metadata = {
            "author": doc.core_properties.author,
            "title": doc.core_properties.title,
            "subject": doc.core_properties.subject,
            "created": str(doc.core_properties.created),
            "modified": str(doc.core_properties.modified),
        }
        
        content = []
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
            paragraph_data = {
                "text": paragraph.text,
                "style": paragraph.style.name,
            }
            content.append(paragraph_data)
        
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            tables.append(table_data)
        
        full_document_structure = {
            "source": filepath,
            "source_ext": filepath.split('.')[-1],
            "source_id": Base64.encode(filepath),
            "metadata": metadata,
            "content": content,
            "tables": tables,
            "images": images_info,
        }
        
        return full_document_structure